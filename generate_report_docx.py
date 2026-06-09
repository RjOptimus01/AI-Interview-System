import math
import os

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = r"F:\MyImportantData\ai-interview-system\AI_Interview_System_Project_Report_With_Diagrams.docx"
ASSET_DIR = r"F:\MyImportantData\ai-interview-system\report_diagrams"


def ensure_asset_dir():
    os.makedirs(ASSET_DIR, exist_ok=True)


def get_font(size, bold=False):
    candidates = []
    if bold:
        candidates.extend(
            [
                r"C:\Windows\Fonts\timesbd.ttf",
                r"C:\Windows\Fonts\arialbd.ttf",
                r"C:\Windows\Fonts\calibrib.ttf",
            ]
        )
    candidates.extend(
        [
            r"C:\Windows\Fonts\times.ttf",
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
        ]
    )
    for path in candidates:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def rounded_rectangle(draw, box, radius, fill, outline, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered_text(draw, box, text, font, fill="#ffffff"):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=4)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = x1 + ((x2 - x1) - text_w) / 2
    y = y1 + ((y2 - y1) - text_h) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, align="center", spacing=4)


def draw_process(draw, box, text, fill="#2f6fd6", outline="#0f2f73"):
    rounded_rectangle(draw, box, radius=18, fill=fill, outline=outline)
    draw_centered_text(draw, box, text, get_font(22, bold=True))


def draw_terminal(draw, box, text, fill="#3283e6", outline="#0f2f73"):
    draw.ellipse(box, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, get_font(22, bold=True))


def draw_decision(draw, center, width, height, text, fill="#3a7be0", outline="#0f2f73"):
    cx, cy = center
    pts = [(cx, cy - height // 2), (cx + width // 2, cy), (cx, cy + height // 2), (cx - width // 2, cy)]
    draw.polygon(pts, fill=fill, outline=outline)
    draw.line(pts + [pts[0]], fill=outline, width=3)
    draw_centered_text(draw, (cx - width // 2, cy - height // 2, cx + width // 2, cy + height // 2), text, get_font(21, bold=True))


def draw_parallelogram(draw, box, text, fill="#3779dd", outline="#0f2f73"):
    x1, y1, x2, y2 = box
    skew = 35
    pts = [(x1 + skew, y1), (x2, y1), (x2 - skew, y2), (x1, y2)]
    draw.polygon(pts, fill=fill, outline=outline)
    draw.line(pts + [pts[0]], fill=outline, width=3)
    draw_centered_text(draw, box, text, get_font(21, bold=True))


def draw_document(draw, box, text, fill="#3f83e3", outline="#0f2f73"):
    x1, y1, x2, y2 = box
    fold = 24
    pts = [(x1, y1), (x2 - fold, y1), (x2, y1 + fold), (x2, y2), (x1, y2)]
    draw.polygon(pts, fill=fill, outline=outline)
    draw.line(pts + [pts[0]], fill=outline, width=3)
    draw.line([(x2 - fold, y1), (x2 - fold, y1 + fold), (x2, y1 + fold)], fill=outline, width=3)
    draw_centered_text(draw, box, text, get_font(20, bold=True))


def draw_manual_input(draw, box, text, fill="#3478da", outline="#0f2f73"):
    x1, y1, x2, y2 = box
    lift = 24
    pts = [(x1 + 18, y1), (x2, y1), (x2 - 18, y2), (x1, y2)]
    draw.polygon(pts, fill=fill, outline=outline)
    draw.line(pts + [pts[0]], fill=outline, width=3)
    draw_centered_text(draw, box, text, get_font(20, bold=True))


def draw_database(draw, box, text, fill="#3579de", outline="#0f2f73"):
    x1, y1, x2, y2 = box
    ellipse_h = 26
    draw.rectangle((x1, y1 + ellipse_h // 2, x2, y2 - ellipse_h // 2), fill=fill, outline=outline, width=3)
    draw.ellipse((x1, y1, x2, y1 + ellipse_h), fill=fill, outline=outline, width=3)
    draw.ellipse((x1, y2 - ellipse_h, x2, y2), fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, get_font(20, bold=True))


def draw_arrow(draw, start, end, fill="#0f2f73", width=5):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    wing = math.pi / 7
    p1 = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    p2 = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, p1, p2], fill=fill)


def draw_title(image, title):
    draw = ImageDraw.Draw(image)
    draw.multiline_text((60, 28), title, font=get_font(28, bold=True), fill="#0f2f73", spacing=4)


def create_canvas(title):
    image = Image.new("RGB", (1400, 900), "#eef5ff")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((18, 18, 1382, 882), radius=24, outline="#9ab8ee", width=3, fill="#f7fbff")
    draw_title(image, title)
    return image, draw


def save_diagram(image, filename):
    path = os.path.join(ASSET_DIR, filename)
    image.save(path, "PNG")
    return path


def build_diagrams():
    ensure_asset_dir()
    paths = {}

    image, draw = create_canvas("Figure 4.1: Overall System Architecture")
    draw_process(draw, (80, 180, 320, 270), "Frontend\n(React UI)")
    draw_process(draw, (420, 180, 660, 270), "Backend API\n(Node + Express)")
    draw_database(draw, (760, 165, 980, 285), "Question\nDatasets")
    draw_process(draw, (1080, 180, 1320, 270), "AI Services\n(OpenRouter)")
    draw_parallelogram(draw, (180, 470, 430, 560), "Resume Upload")
    draw_process(draw, (530, 470, 790, 560), "Evaluation + ATS\nService Layer")
    draw_database(draw, (930, 455, 1180, 575), "Local Storage\nProgress Data")
    draw_arrow(draw, (320, 225), (420, 225))
    draw_arrow(draw, (660, 225), (760, 225))
    draw_arrow(draw, (980, 225), (1080, 225))
    draw_arrow(draw, (300, 515), (530, 515))
    draw_arrow(draw, (790, 515), (930, 515))
    draw_arrow(draw, (200, 270), (200, 470))
    draw_arrow(draw, (1180, 515), (1180, 285))
    paths["architecture"] = save_diagram(image, "figure_4_1_architecture.png")

    image, draw = create_canvas("Figure 4.2: Resume Processing Flowchart")
    draw_terminal(draw, (90, 110, 300, 200), "Start")
    draw_document(draw, (90, 250, 320, 350), "Upload Resume")
    draw_process(draw, (90, 400, 340, 500), "Detect File Type")
    draw_decision(draw, (220, 610), 270, 140, "Valid Format?")
    draw_process(draw, (460, 250, 730, 350), "Extract Text\n(PDF/DOCX/TXT)")
    draw_process(draw, (460, 430, 730, 530), "Normalize Text\nand Sections")
    draw_process(draw, (870, 250, 1160, 350), "Build Candidate\nProfile")
    draw_database(draw, (870, 430, 1160, 550), "Profile Snapshot")
    draw_terminal(draw, (1080, 650, 1300, 740), "End")
    draw_manual_input(draw, (90, 710, 340, 800), "Show Error")
    draw_arrow(draw, (195, 200), (195, 250))
    draw_arrow(draw, (205, 350), (205, 400))
    draw_arrow(draw, (220, 500), (220, 540))
    draw_arrow(draw, (355, 610), (460, 610))
    draw_arrow(draw, (595, 350), (595, 430))
    draw_arrow(draw, (730, 300), (870, 300))
    draw_arrow(draw, (730, 480), (870, 480))
    draw_arrow(draw, (1160, 490), (1190, 650))
    draw_arrow(draw, (220, 680), (220, 710))
    draw.multiline_text((270, 535), "Yes", font=get_font(18, bold=True), fill="#0f2f73")
    draw.multiline_text((150, 685), "No", font=get_font(18, bold=True), fill="#0f2f73")
    paths["resume_flow"] = save_diagram(image, "figure_4_2_resume_flow.png")

    image, draw = create_canvas("Figure 4.3: Interview Question Generation Flowchart")
    draw_terminal(draw, (90, 120, 300, 210), "Start")
    draw_database(draw, (90, 280, 320, 400), "Candidate\nProfile")
    draw_process(draw, (90, 470, 350, 570), "Extract Skills,\nProjects, Keywords")
    draw_database(draw, (500, 120, 760, 240), "Question Bank\nCSV + Parquet")
    draw_process(draw, (500, 330, 790, 430), "Rank Questions by\nRelevance")
    draw_decision(draw, (955, 385), 260, 140, "AI Available?")
    draw_process(draw, (1080, 180, 1320, 280), "Enhance with\nAI Generation")
    draw_process(draw, (1080, 470, 1320, 570), "Finalize 10\nQuestions")
    draw_terminal(draw, (1080, 700, 1300, 790), "End")
    draw_arrow(draw, (195, 210), (195, 280))
    draw_arrow(draw, (205, 400), (205, 470))
    draw_arrow(draw, (350, 520), (500, 380))
    draw_arrow(draw, (630, 240), (630, 330))
    draw_arrow(draw, (790, 380), (825, 380))
    draw_arrow(draw, (1085, 350), (1200, 280))
    draw_arrow(draw, (1085, 420), (1200, 470))
    draw_arrow(draw, (1200, 570), (1200, 700))
    draw.multiline_text((1010, 300), "Yes", font=get_font(18, bold=True), fill="#0f2f73")
    draw.multiline_text((1005, 445), "No", font=get_font(18, bold=True), fill="#0f2f73")
    paths["question_flow"] = save_diagram(image, "figure_4_3_question_flow.png")

    image, draw = create_canvas("Figure 4.4: Answer Evaluation and ATS Analysis")
    draw_terminal(draw, (70, 120, 280, 210), "Start")
    draw_manual_input(draw, (60, 280, 310, 370), "User Answer\nTranscript")
    draw_process(draw, (60, 450, 330, 550), "Compute Quality,\nCoverage, Structure")
    draw_database(draw, (450, 120, 700, 240), "Reference Answers\nand Keywords")
    draw_process(draw, (450, 320, 730, 420), "Score Each Answer")
    draw_process(draw, (450, 520, 730, 620), "Build Overall\nInterview Report")
    draw_document(draw, (860, 120, 1130, 220), "Resume Text")
    draw_process(draw, (860, 310, 1140, 410), "ATS Scoring,\nMistake Detection")
    draw_process(draw, (860, 500, 1140, 600), "Merge Results into\nDashboard")
    draw_terminal(draw, (1060, 700, 1280, 790), "End")
    draw_arrow(draw, (175, 210), (175, 280))
    draw_arrow(draw, (185, 370), (185, 450))
    draw_arrow(draw, (330, 500), (450, 370))
    draw_arrow(draw, (575, 240), (575, 320))
    draw_arrow(draw, (590, 420), (590, 520))
    draw_arrow(draw, (995, 220), (995, 310))
    draw_arrow(draw, (1000, 410), (1000, 500))
    draw_arrow(draw, (730, 570), (860, 550))
    draw_arrow(draw, (1140, 600), (1170, 700))
    paths["eval_ats_flow"] = save_diagram(image, "figure_4_4_eval_ats_flow.png")

    image, draw = create_canvas("Figure 4.5: Camera Monitoring Flowchart")
    draw_terminal(draw, (90, 100, 300, 190), "Start")
    draw_manual_input(draw, (90, 240, 350, 330), "Request Camera\nPermission")
    draw_decision(draw, (220, 450), 270, 140, "Permission\nGranted?")
    draw_process(draw, (460, 240, 730, 340), "Initialize Video\nand Landmarker")
    draw_process(draw, (460, 420, 730, 520), "Track Face Landmarks")
    draw_decision(draw, (960, 470), 280, 150, "Face Visible and\nAttentive?")
    draw_process(draw, (1080, 220, 1320, 320), "Continue Session")
    draw_process(draw, (1080, 580, 1320, 680), "Raise Warning and\nUpdate Metrics")
    draw_terminal(draw, (1080, 760, 1300, 850), "End")
    draw_manual_input(draw, (90, 650, 340, 740), "Show Device Error")
    draw_arrow(draw, (195, 190), (195, 240))
    draw_arrow(draw, (220, 330), (220, 380))
    draw_arrow(draw, (355, 450), (460, 450))
    draw_arrow(draw, (595, 340), (595, 420))
    draw_arrow(draw, (730, 470), (820, 470))
    draw_arrow(draw, (1095, 430), (1200, 320))
    draw_arrow(draw, (1105, 520), (1200, 580))
    draw_arrow(draw, (1200, 680), (1200, 760))
    draw_arrow(draw, (220, 520), (220, 650))
    draw.multiline_text((270, 385), "Yes", font=get_font(18, bold=True), fill="#0f2f73")
    draw.multiline_text((150, 575), "No", font=get_font(18, bold=True), fill="#0f2f73")
    draw.multiline_text((1015, 380), "Yes", font=get_font(18, bold=True), fill="#0f2f73")
    draw.multiline_text((1010, 540), "No", font=get_font(18, bold=True), fill="#0f2f73")
    paths["camera_flow"] = save_diagram(image, "figure_4_5_camera_flow.png")

    return paths


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def base_style(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    document.styles["List Bullet"].font.name = "Times New Roman"
    document.styles["List Bullet"].font.size = Pt(12)


def add_paragraph(document, text, first_line_indent=True, space_after=6):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_heading(document, text, level=1, center=False):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(16)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_table(document, title, headers, rows):
    add_paragraph(document, title, first_line_indent=False, space_after=3)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_paragraph()


def add_bullet_list(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)


def add_diagram_page(document, title, image_path, caption):
    document.add_page_break()
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.line_spacing = 1.2
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    img_p = document.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_after = Pt(6)
    img_p.add_run().add_picture(image_path, width=Inches(6.6))
    cap_p = document.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.line_spacing = 1.2
    cap_run = cap_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.name = "Times New Roman"
    cap_run.font.size = Pt(11)


doc = Document()
base_style(doc)
diagram_paths = build_diagrams()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1)

# Cover page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
r = p.add_run("PROJECT REPORT")
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
r = p.add_run("\nAI INTERVIEW SYSTEM\n")
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.add_run("Submitted in partial fulfilment of the requirement for the award of the degree of\n").font.name = "Times New Roman"
r = p.add_run("BACHELOR OF TECHNOLOGY")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.add_run("\nIN\nCOMPUTER SCIENCE AND ENGINEERING\n").font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.add_run("\nSubmitted by\n[Student Name 1]    University Roll No.: [Roll Number 1]\n[Student Name 2]    University Roll No.: [Roll Number 2]\n").font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.add_run("\nUnder the guidance of\n[Guide Name]\n[Guide Designation]\n").font.name = "Times New Roman"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.add_run("\nDepartment of Computer Science and Engineering\n[University Name]\n[Submission Month, Year]").font.name = "Times New Roman"

doc.add_page_break()

add_heading(doc, "CANDIDATE'S DECLARATION", 1, center=True)
decl = (
    "I/We hereby declare that the project report entitled \"AI Interview System\" submitted in partial fulfillment "
    "of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering "
    "is an original work carried out by us under the guidance of [Guide Name], [Guide Designation], Department of "
    "Computer Science and Engineering, [University Name]. The work presented in this report has not been submitted "
    "to any other university or institution for the award of any degree, diploma, or certificate. All the sources "
    "consulted during the development of the project and the preparation of this report have been properly acknowledged."
)
add_paragraph(doc, decl, first_line_indent=False)
add_paragraph(doc, "[Student Name 1]\nUniversity Roll No.: [Roll Number 1]\nSignature: __________", first_line_indent=False)
add_paragraph(doc, "[Student Name 2]\nUniversity Roll No.: [Roll Number 2]\nSignature: __________", first_line_indent=False)

doc.add_page_break()
add_heading(doc, "ACKNOWLEDGEMENT", 1, center=True)
ack_paras = [
    "The completion of this project has been possible because of the encouragement, guidance, and support extended by many people. We express our sincere gratitude to [Guide Name], [Guide Designation], for valuable guidance, patient mentoring, and continuous motivation throughout the project life cycle. Their suggestions helped us refine the technical direction of the work and encouraged us to maintain both academic rigor and practical relevance.",
    "We are thankful to the Department of Computer Science and Engineering, [University Name], for providing the academic environment necessary to explore software engineering, artificial intelligence, computer vision, and web technologies in an integrated way. The department’s resources, faculty support, and structured learning atmosphere played a major role in the successful realization of this work.",
    "We also acknowledge our classmates, friends, and family members for their constant encouragement and constructive feedback. Their support during implementation, testing, review, and documentation made it easier to complete the project with confidence. We are especially grateful for the motivation that helped us stay consistent during the demanding phases of development and report preparation."
]
for para in ack_paras:
    add_paragraph(doc, para)

doc.add_page_break()
add_heading(doc, "ABSTRACT", 1, center=True)
abstract_paras = [
    "Interview performance has become one of the most critical factors in modern recruitment, especially for students, fresh graduates, and entry-level professionals seeking internships or placements. In spite of having sufficient technical knowledge, many candidates struggle to perform well in interviews because of weak communication, insufficient practice, lack of personalized guidance, and poor awareness of how recruiters and Applicant Tracking Systems evaluate their resumes. Traditional preparation methods generally rely on static question lists, unstructured online resources, and occasional peer-to-peer mock interviews. These approaches often fail to provide role-specific preparation, measurable performance analysis, or immediate corrective feedback. The present project addresses these limitations through the design and implementation of an AI Interview System that combines resume-based personalization, mock interview simulation, ATS analysis, and performance tracking in a unified web application.",
    "The proposed system is designed as a practical interview preparation platform built using React for the frontend and Node.js with Express for the backend. It supports secure user onboarding through login and signup interfaces, followed by resume upload in multiple formats such as PDF, DOCX, and TXT. The uploaded resume is automatically processed using parsing libraries, and the extracted content is used to generate a structured candidate profile consisting of personal details, skills, projects, experience, achievements, and educational information. This profile becomes the foundation for intelligent interview question generation. The system supports HR, Technical, and Combined interview modes so that the user can simulate different placement scenarios in a realistic manner.",
    "A major contribution of the project lies in its answer evaluation mechanism. Users can answer interview questions using voice input, and the application captures or edits the transcript before evaluation. The evaluation pipeline combines dataset-backed heuristics and AI-driven logic to score each answer on dimensions such as relevance, completeness, clarity, speaking quality, and use of examples. It then produces a detailed result sheet containing overall score, rating, strengths, weaknesses, recommendations, and question-wise analysis. This makes the system more useful than a simple quiz interface because it explains where the candidate performed well and where improvement is required.",
    "The system also includes an ATS analysis module to help candidates improve the quality of their resumes before the actual recruitment process. The ATS checker evaluates formatting, readability, section completeness, skill coverage, keyword alignment, action verbs, and common mistakes. This feature extends the project beyond mock interviews and turns it into a broader employability support system. In addition, the system offers an optional camera mode that uses MediaPipe Face Landmarker technology to detect whether the candidate remains visible and attentive during the session. This adds a behavioral dimension to interview practice and helps simulate real-world discipline.",
    "To maintain user engagement over time, the project introduces gamification features such as experience points, ranks, streaks, badges, and daily challenge logic. Session history and progress records are also preserved locally, enabling users to track long-term improvement. Overall, the AI Interview System demonstrates how artificial intelligence, document processing, web development, and computer vision can be integrated to create a practical and extensible interview preparation platform. The final outcome is a technically valid and user-centered system that can significantly enhance preparation quality while remaining accessible, modular, and suitable for further enhancement."
]
for para in abstract_paras:
    add_paragraph(doc, para)

doc.add_page_break()
add_heading(doc, "TABLE OF CONTENTS", 1, center=True)
toc_lines = [
    "Candidate's Declaration",
    "Acknowledgement",
    "Abstract",
    "List of Tables",
    "List of Figures",
    "Abbreviations",
    "Notations",
    "Chapter 1: Introduction and Motivation",
    "Chapter 2: Objectives or Problem Statement",
    "Chapter 3: Software and Hardware Requirements",
    "Chapter 4: Project Methodology / Design",
    "Chapter 5: Implementation and Results",
    "Chapter 6: Conclusion and Future Scope",
    "References",
    "Appendix A: Source Code Overview",
    "Appendix B: Project Screenshots",
]
for line in toc_lines:
    add_paragraph(doc, line, first_line_indent=False, space_after=0)

add_heading(doc, "LIST OF TABLES", 1, center=True)
tables_list = [
    "Table 3.1 Software Requirements",
    "Table 3.2 Hardware Requirements",
    "Table 3.3 Functional Requirements",
    "Table 3.4 Non-Functional Requirements",
    "Table 4.1 Major System Modules",
    "Table 5.1 API Endpoints of the System",
    "Table 5.2 Evaluation Parameters",
    "Table 5.3 ATS Analysis Parameters",
]
for line in tables_list:
    add_paragraph(doc, line, first_line_indent=False, space_after=0)

add_heading(doc, "LIST OF FIGURES", 1, center=True)
fig_list = [
    "Figure 4.1 Overall System Architecture",
    "Figure 4.2 Use Case Diagram of AI Interview System",
    "Figure 4.3 Resume Processing Workflow",
    "Figure 4.4 Interview Question Generation Flow",
    "Figure 4.5 Answer Evaluation Flow",
    "Figure 4.6 ATS Analysis Workflow",
    "Figure 4.7 Camera Monitoring Workflow",
    "Figure 5.1 Login Page",
    "Figure 5.2 Resume Upload Interface",
    "Figure 5.3 Interview Session Screen",
    "Figure 5.4 ATS Analysis Dashboard",
    "Figure 5.5 Result Summary Screen",
]
for line in fig_list:
    add_paragraph(doc, line, first_line_indent=False, space_after=0)

add_heading(doc, "ABBREVIATIONS", 1, center=True)
abbrs = [
    "AI - Artificial Intelligence",
    "ATS - Applicant Tracking System",
    "API - Application Programming Interface",
    "UI - User Interface",
    "UX - User Experience",
    "LLM - Large Language Model",
    "PDF - Portable Document Format",
    "DOCX - Word Document Format",
    "CSV - Comma-Separated Values",
    "XP - Experience Points",
    "SRS - Software Requirement Specification",
    "HR - Human Resources",
]
for line in abbrs:
    add_paragraph(doc, line, first_line_indent=False, space_after=0)

add_heading(doc, "NOTATIONS", 1, center=True)
notations = [
    "q - Interview question instance",
    "a - Candidate answer instance",
    "R - Resume text input",
    "S - Answer score",
    "K - Extracted keyword set",
    "M - Camera monitoring state",
]
for line in notations:
    add_paragraph(doc, line, first_line_indent=False, space_after=0)

doc.add_page_break()

# Chapter 1
add_heading(doc, "CHAPTER 1", 1, center=True)
add_heading(doc, "INTRODUCTION AND MOTIVATION", 1, center=True)
chapter1 = {
    "1.1 Background of Interview Preparation": [
        "The interview process has evolved from being a short face-to-face assessment into a multidimensional evaluation of technical competence, communication ability, behavioral readiness, and professional maturity. In campus placements and early-career recruitment, employers increasingly expect candidates to demonstrate not only knowledge but also the ability to explain projects clearly, justify technical choices, communicate confidently, and adapt to different categories of questions. This has made interview preparation a skill-intensive activity rather than a purely knowledge-based exercise.",
        "Many students enter interviews after studying theory, solving coding problems, or reading frequently asked questions, yet they still feel underprepared when the conversation becomes personalized. A candidate may know programming concepts but fail to explain the architecture of a project, the reasoning behind a design choice, or the impact of a team contribution. This gap between technical ability and interview delivery is one of the major reasons why otherwise capable candidates underperform in selection processes.",
        "The digital transformation of recruitment has also changed the preparation landscape. Resumes are now often screened by Applicant Tracking Systems before a human interviewer even reviews them. As a result, candidates need support at two levels: resume readiness and interview readiness. A preparation platform that addresses only one of these concerns remains incomplete. This project is therefore positioned in the broader context of employability support, where the goal is to help candidates present their capabilities more effectively from the screening stage through to the interview stage."
    ],
    "1.2 Problems in Traditional Interview Practice": [
        "Traditional preparation methods typically include browsing websites for common interview questions, practicing with friends, joining occasional mock interviews, or watching tutorial videos. Although these approaches are useful at an introductory level, they have several limitations. The first limitation is lack of personalization. Most available materials are generic and do not account for the candidate's own resume, projects, internships, or technical specialization. This creates a mismatch between preparation content and likely interview experience.",
        "The second limitation is the lack of immediate and structured feedback. In a real interview, quality is judged not only by whether an answer contains correct information but also by how clearly it is delivered, how well it addresses the question, and whether it includes relevant examples. Traditional practice methods often fail to capture these dimensions. Candidates may answer a question and move on without knowing whether they were concise, relevant, complete, or impressive.",
        "A third problem is inconsistency. Manual mock interviews depend on the availability, mood, and quality of the evaluator. Two different evaluators may offer very different suggestions for the same answer. Similarly, some students do not have access to mentors or seniors who can regularly conduct interview practice sessions. This creates an uneven preparation experience. An intelligent software system can provide repeatable practice sessions at any time and can maintain a consistent evaluation framework."
    ],
    "1.3 Need for an AI-Driven Interview System": [
        "Artificial intelligence provides the means to create adaptive systems that are more aligned with real user needs. In the context of interview preparation, AI can interpret resume content, identify candidate strengths, generate role-aware questions, and analyze the textual quality of responses. Such a system does not replace human judgment but enhances preparation by providing rapid, structured, and personalized feedback in a scalable manner.",
        "An AI-driven interview system is especially useful because it can bridge the gap between static preparation resources and dynamic interview interaction. By analyzing skills, projects, tools, and achievements from the candidate's resume, the system can tailor the mock interview to the actual background of the user. This increases relevance and makes practice more meaningful. It also encourages users to understand their own resume better, which is essential for success in technical and HR interviews.",
        "The combination of AI evaluation and dataset-backed logic further improves reliability. Instead of relying entirely on generative text, the system can compare answers against expected topic coverage, question intent, and known focus areas. This produces richer analysis such as missing points, examples used, transcript quality, and recommended improvement direction. The result is a preparation environment that feels more practical, measurable, and supportive."
    ],
    "1.4 Motivation of the Project": [
        "The motivation behind this project comes from the recurring difficulty observed among students during placement preparation. Many candidates are anxious about interviews because they do not know what type of questions they will face, whether their answers are convincing, or how their resume will be perceived. In addition, there is often no single platform that combines resume review, interview simulation, answer evaluation, and progress tracking in one connected workflow.",
        "Another motivation is the need for practice that goes beyond text-based rehearsal. Real interviews require voice delivery, confidence, focus, and discipline. By including voice-based response input and optional camera monitoring, the project attempts to simulate aspects of real interview pressure in a controlled and supportive setting. The system also uses gamification to maintain motivation over repeated sessions, which is essential because interview readiness improves through consistent practice rather than one-time effort.",
        "From a technical learning perspective, the project offers a valuable opportunity to combine frontend engineering, backend API design, AI service integration, document parsing, heuristic scoring, dataset handling, and computer vision. Building such a system demonstrates how multiple domains of computer science can work together to solve a practical problem with clear user value."
    ],
    "1.5 Scope of the Proposed System": [
        "The AI Interview System is designed as a web-based mock interview platform targeted primarily at students, fresh graduates, and early-career job seekers. The system supports resume upload in commonly used formats, extracts structured candidate information, and uses that information to drive multiple interview preparation workflows. Users can attempt HR interviews, technical interviews, or combined sessions that simulate mixed rounds.",
        "The project scope includes intelligent question generation, answer capture using voice or manual editing, answer evaluation, ATS score generation, camera-based attentiveness monitoring, and gamified progress tracking. The system stores user-related information such as recent sessions, rewards, and practice history locally so that candidates can observe improvement trends over time.",
        "At the same time, the present version is not intended to act as a recruitment decision engine. It is a preparation support tool rather than an automated hiring platform. Features such as recruiter-side analytics, cloud deployment, large-scale multi-user databases, or real-time speech emotion scoring are considered future scope rather than current deliverables."
    ],
    "1.6 Objectives Overview": [
        "The central objective of the project is to create a practical and intelligent interview practice environment that gives users profile-specific preparation rather than one-size-fits-all question lists. The system must reduce the distance between resume content and mock interview behavior by ensuring that generated questions reflect the actual background of the candidate.",
        "A parallel objective is to provide structured performance feedback through scores, strengths, weaknesses, and recommendations. The user should be able to understand why an answer was strong or weak and what changes would improve performance in the next session. This turns the platform into a training tool rather than a passive question display system.",
        "The broader objective is to integrate employability support features into one cohesive application. Resume parsing, ATS analysis, answer evaluation, camera monitoring, and gamification together form a more complete preparation ecosystem. In this sense, the project aims not only to simulate an interview but also to improve the overall quality of preparation."
    ]
}
for heading, paras in chapter1.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

doc.add_page_break()

# Chapter 2
add_heading(doc, "CHAPTER 2", 1, center=True)
add_heading(doc, "OBJECTIVES OR PROBLEM STATEMENT", 1, center=True)
chapter2 = {
    "2.1 Problem Statement": [
        "Candidates preparing for placements often rely on fragmented tools and generic practice resources. One platform may offer common interview questions, another may provide a resume score, and a different tool may support coding practice. This separation makes preparation inefficient and prevents users from viewing interview readiness as a connected process. The absence of personalization leads to shallow preparation, while the absence of analytics makes self-improvement difficult.",
        "The specific problem addressed in this project is the lack of an accessible and integrated system that can read a candidate's resume, generate relevant questions, capture interview responses, evaluate answer quality, analyze ATS readiness, and maintain user progress over time. Without such a system, candidates must piece together different preparation methods and still remain uncertain about their actual readiness."
    ],
    "2.2 Issues Faced by Students and Job Seekers": [
        "Students commonly face uncertainty regarding what interviewers expect from them. Technical questions may be based on core subjects, current projects, or tools mentioned in the resume, while HR questions often require reflective and situation-based answers. When students practice only generic questions, they remain unprepared for resume-specific probing.",
        "Another major issue is lack of objective feedback. Candidates may not know whether an answer was too brief, poorly structured, unsupported by examples, or weak in technical depth. Manual feedback is often irregular, while self-evaluation is usually inaccurate. Additionally, many candidates are unaware that their resume itself may be limiting their chances because of missing keywords, poor formatting, or weak action language.",
        "The issue of consistency is also important. Practice tends to happen irregularly because students lose motivation or fail to see measurable growth. A system that records progress, rewards repetition, and highlights improvement patterns can help address this problem."
    ],
    "2.3 Main Objectives": [
        "The first objective of the project is to design and implement a web application that can automate personalized interview practice using the candidate's uploaded resume. The system should generate relevant interview questions and support more than one interview mode so that users can simulate different recruitment scenarios.",
        "The second objective is to build an evaluation pipeline that analyzes the user's responses and produces meaningful output. This output should include question-wise score, overall score, answer quality indicators, strengths, weaknesses, and improvement suggestions. The result must be understandable to students and actionable enough to guide future preparation.",
        "The third objective is to integrate ATS analysis into the same platform. Since resume quality influences interview opportunities, the system should help users detect resume-related problems such as missing keywords, incomplete sections, or formatting issues before they apply for jobs."
    ],
    "2.4 Secondary Objectives": [
        "The project also pursues secondary objectives related to realism, usability, and engagement. One secondary objective is to support voice-based answering so that users practice speaking, not just typing. Another is to provide optional camera monitoring that checks whether the user's face remains visible and attentive during the interview session.",
        "An additional objective is to keep users motivated through gamification elements such as XP, streaks, rank progression, and daily challenges. These features encourage repeated usage and help convert interview practice into an ongoing habit.",
        "A final secondary objective is modularity. The design should allow new features to be added in the future, such as multilingual interviews, recruiter dashboards, cloud data storage, emotion analysis, or more advanced speech scoring mechanisms."
    ],
    "2.5 Expected Outcomes": [
        "The expected output of the project is a functioning AI Interview System that gives users a more realistic and informative preparation experience than static question banks. Users should be able to upload a resume, select an interview type, answer generated questions, receive performance analysis, and improve iteratively.",
        "From the ATS perspective, the expected outcome is that users gain better awareness of how resume structure and wording affect screening readiness. The system should make resume weaknesses explicit enough that users can edit them confidently.",
        "From an academic standpoint, the project is expected to demonstrate a meaningful integration of web technologies, AI services, document parsing, heuristic scoring, dataset handling, and client-side interaction patterns. It should therefore contribute both as a practical product and as a technical study."
    ],
    "2.6 Advantages of the Proposed System": [
        "The proposed system offers several clear advantages. It is personalized because it uses resume content to generate relevant questions. It is immediate because it delivers feedback at the end of each session without requiring human evaluation. It is comprehensive because it combines interview simulation, resume analysis, camera attention support, and progress tracking in one application.",
        "Another advantage is consistency. The same logic is applied to each session, which gives users a stable evaluation framework. The platform is also accessible because it can be used repeatedly on demand, without requiring an interviewer to be available. Finally, its modular architecture makes it suitable for future growth in both academic and practical settings."
    ]
}
for heading, paras in chapter2.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

doc.add_page_break()

# Chapter 3
add_heading(doc, "CHAPTER 3", 1, center=True)
add_heading(doc, "SOFTWARE AND HARDWARE REQUIREMENTS", 1, center=True)
chapter3 = {
    "3.1 Introduction": [
        "A well-defined requirement analysis is essential for the successful development of any software system. It helps identify what the system must do, what quality standards it should meet, which resources are needed, and what practical constraints govern its operation. For the AI Interview System, requirement analysis is especially important because the application combines document processing, AI integration, real-time user interaction, camera access, and performance evaluation in a single workflow.",
        "This chapter presents the functional requirements, non-functional requirements, software and hardware needs, and assumptions associated with the implementation of the project. It also describes the expected operating environment and gives a concise software requirement specification suitable for an academic project report."
    ],
    "3.2 Functional Requirements": [
        "The system shall allow the user to create an account and log in securely through a frontend authentication interface. The system shall store and retrieve user details for future sessions.",
        "The system shall allow resume upload in PDF, DOCX, and TXT formats. It shall extract text from the uploaded file and derive a structured candidate profile containing fields such as name, email, skills, projects, achievements, and education.",
        "The system shall generate interview questions based on resume content and allow the user to choose among HR, Technical, and Combined interview types. It shall support answer capture through voice transcription and manual text review, evaluate completed interview sessions, and present a detailed result summary.",
        "The system shall provide ATS score analysis for the resume and display strengths, mistakes, keyword matching information, and improvement priorities. It shall also support optional camera monitoring during interview sessions and display warning information based on attention conditions. Finally, the system shall preserve session history and gamification-related progress data."
    ],
    "3.3 Non-Functional Requirements": [
        "The system should provide a responsive and intuitive user interface so that users with limited technical experience can navigate the workflow easily. Pages must remain readable and organized on standard laptop screens used for placement preparation.",
        "The application should process normal resumes within a practical time window and return interview questions or ATS analysis without unreasonable delay. Since some actions involve AI calls and parsing logic, graceful loading indicators and informative error handling are required.",
        "The system should remain modular and maintainable. Frontend responsibilities, backend routes, and service logic should be separated so that future modifications can be made without rewriting the entire application. Reliability, clarity of feedback, and usability are therefore treated as major non-functional expectations."
    ],
    "3.4 Software Requirement Specification": [
        "At the software level, the frontend requires a JavaScript runtime and browser environment capable of handling modern React applications, media devices, and local storage. The backend requires Node.js, Express, middleware support for uploads, and libraries for document parsing and dataset processing. The project also depends on external AI services through OpenRouter for profile enhancement, question generation, and evaluation support.",
        "The software environment must include React, React DOM, Express, Multer, CORS, Mammoth, PDF parsing utilities, hyparquet or equivalent data access support, and MediaPipe tasks for face landmark detection. Since the project is built for development and demonstration, the runtime environment assumes internet access for AI integration and model loading where applicable."
    ],
    "3.5 Hardware Requirements": [
        "The hardware requirements are moderate and aligned with common student laptops or desktop machines. A system with at least a dual-core processor, 4 GB of RAM, 1 GB of available storage for dependencies and project files, a microphone, and an optional webcam is sufficient for development and usage of the current system.",
        "Camera monitoring features require access to a working webcam, while voice-based input requires a functional microphone and browser support for speech recognition. For smoother development and testing, a system with 8 GB RAM and a stable internet connection is preferred."
    ],
    "3.6 User Requirements": [
        "The primary user of the system is a student or job seeker who wants to improve interview performance. Such a user should be able to sign in, upload a resume, choose interview mode, attempt a mock session, review feedback, and revisit past progress without external assistance.",
        "Another user expectation is understandable feedback. The user does not need raw model output; instead, the system must convert analysis into clear strengths, weaknesses, scores, and suggestions. The user also expects the application to be forgiving of normal issues such as minor transcript errors or short-term camera interruptions."
    ],
    "3.7 Constraints and Assumptions": [
        "The system currently assumes English-language resume content and English interview flow. It assumes that the uploaded resume contains enough extractable text to identify meaningful candidate details. The voice workflow depends on browser and permission support, and the camera workflow depends on hardware availability and permission approval.",
        "Another practical constraint is that AI-powered behavior depends on external API availability and correct environment configuration. In the absence of AI responses, the system is expected to fall back to dataset-backed or heuristic logic wherever possible. These assumptions are acceptable for the current academic scope."
    ]
}
for heading, paras in chapter3.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

add_table(
    doc,
    "Table 3.1 Software Requirements",
    ["Component", "Requirement"],
    [
        ["Frontend", "React, JavaScript, CSS/Tailwind, browser support"],
        ["Backend", "Node.js, Express"],
        ["Upload Handling", "Multer"],
        ["Document Parsing", "Mammoth, PDF parsing library"],
        ["AI Integration", "OpenRouter API"],
        ["Dataset Support", "CSV and Parquet readers"],
        ["Camera Monitoring", "MediaPipe Face Landmarker"],
    ],
)

add_table(
    doc,
    "Table 3.2 Hardware Requirements",
    ["Hardware", "Minimum Requirement"],
    [
        ["Processor", "Dual-core CPU"],
        ["RAM", "4 GB minimum, 8 GB recommended"],
        ["Storage", "1 GB free space for project and dependencies"],
        ["Microphone", "Required for voice interview mode"],
        ["Webcam", "Required for camera interview mode"],
        ["Internet", "Required for AI service access and model resources"],
    ],
)

add_table(
    doc,
    "Table 3.3 Functional Requirements",
    ["ID", "Description"],
    [
        ["FR1", "User signup and login"],
        ["FR2", "Resume upload and text extraction"],
        ["FR3", "Candidate profile generation"],
        ["FR4", "Resume-based interview question generation"],
        ["FR5", "Interview answer capture and evaluation"],
        ["FR6", "ATS score and mistake analysis"],
        ["FR7", "Gamification and history tracking"],
    ],
)

add_table(
    doc,
    "Table 3.4 Non-Functional Requirements",
    ["Category", "Expectation"],
    [
        ["Usability", "Simple, understandable, and guided workflow"],
        ["Performance", "Reasonable response time with visible loading states"],
        ["Maintainability", "Modular frontend and backend organization"],
        ["Reliability", "Graceful handling of API or permission failures"],
        ["Scalability", "Extensible architecture for future features"],
    ],
)

doc.add_page_break()

# Chapter 4
add_heading(doc, "CHAPTER 4", 1, center=True)
add_heading(doc, "PROJECT METHODOLOGY / DESIGN", 1, center=True)
chapter4 = {
    "4.1 Development Methodology": [
        "The system was developed using an incremental and modular methodology. Instead of building the entire platform as one monolithic feature set, the application was organized into independent but connected modules such as authentication, resume processing, interview generation, evaluation, ATS analysis, and gamification. This approach made it easier to implement, test, and improve each unit while preserving overall coherence.",
        "The methodology combined practical software engineering with rapid validation. Core flows were implemented first, followed by feature enrichment and UI refinement. This is especially suitable for an academic project because it balances feasibility and technical depth while allowing progressive integration of AI and computer vision capabilities."
    ],
    "4.2 Overall System Architecture": [
        "The architecture of the AI Interview System follows a client-server model. The React frontend is responsible for user interaction, navigation, form handling, voice input flow, camera UI, and result display. The Node.js and Express backend is responsible for processing uploaded resumes, generating interview questions, evaluating answers, and computing ATS analysis. The system also uses local storage on the client side for persistence of user-related information and practice progress.",
        "The backend service layer contains the major business logic. It loads and normalizes data from CSV and Parquet datasets, extracts text from uploaded documents, derives candidate information, selects or generates questions, and computes evaluation results. AI service calls are encapsulated within the backend so that the frontend receives already-structured JSON responses. This separation improves maintainability and keeps presentation concerns away from core analysis logic.",
    ],
    "4.3 Module-Wise Design": [
        "The authentication module manages user signup and login through local storage-backed data handling. The resume processing module accepts uploads and converts the file into machine-readable text using specialized parsers for DOCX and PDF formats. The question generation module takes candidate profile data and interview type as input, then creates a relevant question set using a blend of datasets and AI-supported generation.",
        "The interview module manages session flow, including question navigation, transcript capture, progress tracking, and optional camera state display. The evaluation module aggregates answer metadata, evaluates transcript quality and content relevance, and returns a question-wise breakdown. The ATS module separately scores the uploaded resume on multiple dimensions. The gamification module updates XP, rank, streak, and challenge status based on completed sessions."
    ],
    "4.4 Resume Parsing Pipeline": [
        "The resume parsing workflow begins with file upload from the frontend. The file is posted to the backend through a multipart form request. Based on the extension of the uploaded file, the backend selects the appropriate parser. TXT files are read directly, DOCX files are processed with Mammoth, and PDF files are handled using a PDF text extraction library. The resulting raw text is normalized to reduce spacing issues, remove formatting artifacts, and create a clean textual base for further analysis.",
        "After text extraction, the backend derives a structured candidate profile. It searches for name patterns, contact details, education, experience, projects, certifications, achievements, and skills. Section-based heuristics are used to organize the extracted content, while optional AI enhancement may refine the profile if external AI support is available. The final profile snapshot is returned to the frontend and also reused for question generation and ATS analysis.",
    ],
    "4.5 Interview Question Generation Flow": [
        "Question generation is based on both the candidate profile and available question banks. The system first loads dataset-backed questions from CSV and Parquet sources. It then filters, normalizes, and scores these records according to relevance. Keywords extracted from the resume, role hints inferred from projects and skills, and the selected interview type collectively influence the final ranking of candidate questions.",
        "If AI services are available, an additional generation layer can create more contextualized questions. However, the system is designed to remain functional even without that layer by falling back to dataset-driven selection and internally generated question templates. This hybrid approach improves robustness because the user receives a meaningful interview set even if the AI layer is unavailable.",
    ],
    "4.6 Interview Evaluation Flow": [
        "The evaluation process begins after the candidate completes the interview session. For each question, the system receives the transcript, answer duration, and selected metadata such as camera state when applicable. The evaluation layer computes speaking-related indicators, keyword relevance, structural completeness, and alignment with expected answer content. It can use both heuristic logic and AI-assisted evaluation depending on environment readiness.",
        "Question-wise results are merged into an aggregate report. This includes content score, speaking score, overall rating, strengths, weaknesses, improvement suggestions, and missing expected points. The system also computes session-level overall score and textual feedback so that the user can understand high-level performance without reading every detail. This layered evaluation strategy is both interpretable and extensible.",
    ],
    "4.7 ATS Analysis Flow": [
        "The ATS analysis module processes the same resume text with a different objective. Instead of generating interview content, it evaluates how readable, complete, and recruiter-friendly the resume appears from an ATS perspective. The module analyzes formatting simplicity, section availability, keyword match, action verbs, weak phrasing, quantified impact indicators, and job-description alignment when a target role description is provided.",
        "The result is organized into category scores, a summary, strengths, detected mistakes, keyword analysis, and improvement priorities. This design is useful because candidates often need specific and prioritized suggestions rather than abstract recommendations. The ATS workflow therefore complements the interview evaluation flow and gives the user an earlier checkpoint in the recruitment pipeline.",
    ],
    "4.8 Camera Monitoring Workflow": [
        "The optional camera mode uses MediaPipe Face Landmarker to estimate whether the candidate remains visible and oriented toward the screen during the session. After user permission is granted, the webcam stream is initialized on the frontend. Face landmarks are then processed frame by frame, and rules related to nose offset, eye tilt, face visibility, and sustained absence are used to estimate attentiveness.",
        "The monitoring output is not intended as a disciplinary feature in the current version. Instead, it acts as informational feedback that encourages the user to maintain a more realistic interview posture. Summary metrics such as attentive ratio, visible-face ratio, warning count, and session mode are included in the final report. This provides a lightweight but useful behavioral signal.",
    ],
    "4.9 Use Case and Activity View": [
        "From a use case perspective, the primary actor is the candidate. The candidate creates an account, uploads a resume, chooses interview type, attempts an interview, checks ATS score, reviews results, and tracks progress. Supporting system actors include AI services and dataset resources. This use case design emphasizes the candidate-centered nature of the platform.",
        "The activity flow begins with authentication, continues into resume acquisition, profile extraction, and branching into either ATS analysis or interview preparation. The interview branch then moves through question generation, answering, evaluation, and result review. Both branches can update stored user progress. This flow captures the practical usage cycle of the application."
    ],
    "4.10 Algorithms and Pseudocode": [
        "The following subsections summarize the high-level algorithms used by the system."
    ],
}
for heading, paras in chapter4.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

add_table(
    doc,
    "Table 4.1 Major System Modules",
    ["Module", "Purpose"],
    [
        ["Authentication", "Manage user login, signup, and session identity"],
        ["Resume Parser", "Extract text and structured profile data from resume"],
        ["Question Generator", "Produce HR, technical, or combined question sets"],
        ["Interview Engine", "Drive question flow and answer capture"],
        ["Evaluation Engine", "Score answers and produce recommendations"],
        ["ATS Analyzer", "Evaluate resume for screening readiness"],
        ["Gamification", "Track XP, streaks, and engagement"],
    ],
)

for algo_title, algo_lines in [
    ("4.10.1 Algorithm for Resume Extraction", [
        "Step 1: Accept uploaded file and validate extension.",
        "Step 2: Select parser based on file type.",
        "Step 3: Extract raw text and normalize whitespace.",
        "Step 4: Detect sections such as skills, projects, and experience.",
        "Step 5: Build candidate profile object and return result.",
    ]),
    ("4.10.2 Algorithm for Interview Question Generation", [
        "Step 1: Load candidate profile and interview type.",
        "Step 2: Extract keywords and role hints from the profile.",
        "Step 3: Rank dataset questions based on relevance.",
        "Step 4: Optionally enhance with AI-generated questions.",
        "Step 5: Return a balanced set of final questions.",
    ]),
    ("4.10.3 Algorithm for AI Answer Evaluation", [
        "Step 1: Read transcript, duration, and question context.",
        "Step 2: Compute structure, relevance, and keyword coverage.",
        "Step 3: Estimate speaking quality using transcript features.",
        "Step 4: Merge heuristic and AI evaluation outputs.",
        "Step 5: Generate question-wise and overall feedback report.",
    ]),
    ("4.10.4 Algorithm for ATS Scoring", [
        "Step 1: Read normalized resume text.",
        "Step 2: Check formatting and section completeness.",
        "Step 3: Compare keywords against expected terms or job description.",
        "Step 4: Detect mistakes and compute weighted category scores.",
        "Step 5: Return overall ATS score and improvement priorities.",
    ]),
    ("4.10.5 Algorithm for Camera Attention Monitoring", [
        "Step 1: Initialize webcam and face landmark model.",
        "Step 2: Process frames and detect face visibility.",
        "Step 3: Measure pose indicators such as tilt and offset.",
        "Step 4: Trigger warnings for sustained inattentive states.",
        "Step 5: Summarize visibility and attention ratios for final report.",
    ]),
]:
    add_heading(doc, algo_title, 3)
    add_bullet_list(doc, algo_lines)

add_diagram_page(doc, "SYSTEM ARCHITECTURE DIAGRAM", diagram_paths["architecture"], "Figure 4.1: Overall System Architecture")
add_diagram_page(doc, "RESUME PROCESSING FLOWCHART", diagram_paths["resume_flow"], "Figure 4.2: Resume Processing Flowchart")
add_diagram_page(doc, "QUESTION GENERATION FLOWCHART", diagram_paths["question_flow"], "Figure 4.3: Interview Question Generation Flowchart")
add_diagram_page(doc, "EVALUATION AND ATS FLOWCHART", diagram_paths["eval_ats_flow"], "Figure 4.4: Answer Evaluation and ATS Analysis Flowchart")
add_diagram_page(doc, "CAMERA MONITORING FLOWCHART", diagram_paths["camera_flow"], "Figure 4.5: Camera Monitoring Flowchart")

doc.add_page_break()

# Chapter 5
add_heading(doc, "CHAPTER 5", 1, center=True)
add_heading(doc, "IMPLEMENTATION AND RESULTS", 1, center=True)
chapter5 = {
    "5.1 Overview of Implementation": [
        "The implementation of the AI Interview System was carried out in a modular manner. The frontend was developed using React and JavaScript, with the interface organized into dedicated components for authentication, home flow, interview practice, dashboards, and shared UI elements. The backend was implemented using Node.js and Express, exposing clearly separated routes for resume upload, question generation, interview evaluation, and ATS analysis. This separation simplified debugging and improved maintainability.",
        "The system was designed to deliver a smooth user journey from registration through result review. Each user action on the frontend maps to a corresponding backend workflow. The frontend handles user interactions and presentation, while the backend concentrates on parsing, analysis, AI communication, and structured result formation."
    ],
    "5.2 Authentication Module": [
        "The authentication module provides signup and login interfaces and stores user data locally for the current project scope. During signup, the system validates that all fields are filled, that the password confirmation matches, and that a duplicate account does not already exist. During login, stored credentials are retrieved and compared with entered values before access is granted.",
        "Although the current implementation uses local storage rather than a production-grade database and encrypted authentication stack, it is suitable for academic demonstration and internal usage. The modular design also allows future migration to a server-backed authentication service without redesigning the user interface."
    ],
    "5.3 Resume Upload and Parsing Module": [
        "The resume upload module is one of the core features of the system because most downstream processing depends on it. The frontend accepts PDF, DOCX, and TXT files and sends the selected file to the backend through a multipart request. The backend validates file presence, size constraints, and supported format before proceeding with text extraction.",
        "After extracting text, the system identifies likely values for personal details, technical skills, projects, education, certifications, and achievements. It then displays the derived profile to the user. This preview is useful because it confirms whether the system has correctly interpreted the resume and helps users trust the generated interview questions.",
        "[Insert Figure 5.2: Resume Upload Interface]"
    ],
    "5.4 Question Generation Module": [
        "Question generation is performed after resume parsing and interview type selection. The system supports HR, Technical, and Combined modes. For a technical round, the question selector emphasizes categories such as fundamentals, debugging, tools, implementation choices, and project architecture. For an HR round, it emphasizes topics like strengths and weaknesses, teamwork, leadership, communication, adaptability, and motivation.",
        "The question generation logic combines structured datasets with AI support. This hybrid design ensures that the system remains operational even when external AI responses are unavailable. The final question set is displayed in a guided session format where the candidate answers one question at a time."
    ],
    "5.5 Interview Session and Voice Answer Module": [
        "The interview session interface displays the current question, session progress, answer area, and controls for voice input. When the user starts voice recording, the system captures spoken input and appends transcript text into the answer area. Users can also manually edit the answer before moving forward. This design improves usability because speech recognition may occasionally need correction.",
        "The interface also supports replaying the question and moving step-by-step through the interview. By structuring the interview as a guided sequence rather than a free-form page, the system keeps users focused on the current task and better simulates actual interview pacing."
    ],
    "5.6 Camera Mode Implementation": [
        "When the user chooses the camera interview mode, the system performs a compatibility check for browser support, microphone access, and camera permission. Once approved, the webcam preview becomes visible on the interview page. MediaPipe Face Landmarker processes visual input in real time and computes simple attention indicators such as whether the face remains visible and roughly aligned with the screen.",
        "If the system detects sustained absence or inattentive movement, it raises warnings without interrupting the session. This keeps the feature helpful rather than punitive. At the end of the session, summary metrics are added to the report to inform the user how consistently attentiveness was maintained."
    ],
    "5.7 Evaluation Result Generation": [
        "After the interview is completed, the candidate's answers are posted to the backend for evaluation. The evaluation engine computes per-question scores and combines them into an overall rating. It also extracts useful descriptive signals such as whether concrete examples were given, whether the answer covered important points, whether speaking quality appeared weak, and which improvement directions are most relevant.",
        "The frontend then presents a full result view containing overall score, rating, reward summary, camera summary if enabled, strengths, areas for improvement, and question-by-question analysis. This result page is intentionally detailed so that users can treat the session as a learning experience rather than simply viewing a numeric score."
    ],
    "5.8 ATS Checker Implementation": [
        "The ATS checker module is implemented as a distinct workflow that reuses the uploaded resume and candidate profile. The user may optionally provide a job description for more targeted keyword analysis. The backend computes category scores related to formatting, ATS readability, section completeness, keyword match, experience quality, achievement impact, and grammar clarity.",
        "Mistakes are also identified and prioritized. Examples include weak sectioning, missing contact details, insufficient quantified impact, or low alignment with target job keywords. The final ATS report gives candidates an actionable checklist for improving resume quality before applying for internships or jobs."
    ],
    "5.9 Gamification and Dashboard Features": [
        "To encourage continued practice, the application includes a gamification layer that awards experience points after completed sessions, tracks streak progression, assigns ranks, and supports daily challenge logic. These features are useful because interview improvement depends on repeated practice. By rewarding consistency, the system promotes habitual engagement rather than one-time use.",
        "The dashboard area also surfaces historical information such as recent ATS runs, stored interview results, and profile progress. This creates a sense of continuity and helps the user measure growth over time."
    ],
    "5.10 Backend APIs and Data Flow": [
        "The backend exposes a health route and four major interview-related endpoints: resume parsing, question generation, interview evaluation, and ATS scoring. Each route validates request structure, invokes the appropriate service function, and returns JSON output to the frontend. This API-driven design keeps the frontend independent from internal backend implementation details.",
        "The service layer is responsible for text extraction, data normalization, question bank loading, scoring logic, and AI request encapsulation. Since the service layer returns structured objects rather than presentation strings, the frontend remains free to render the same data in different layouts or future views."
    ],
    "5.11 Testing Observations and Result Discussion": [
        "Testing was carried out primarily at functional and integration levels. Frontend components were reviewed for correct navigation, state handling, conditional rendering, and session progression. Backend endpoints were tested for missing inputs, invalid request shapes, and normal usage cases. Resume uploads in supported formats were verified along with API responses for parsing, question generation, evaluation, and ATS scoring.",
        "The overall results indicate that the application is able to produce a cohesive preparation flow. Resume-based question selection adds practical relevance, and the evaluation summaries provide actionable guidance. ATS analysis strengthens the usefulness of the platform by helping users improve before interview day. Camera monitoring and gamification further contribute to a more engaging and realistic experience. While the project remains a demonstration-scale academic system, it successfully proves the viability of integrating document analysis, AI-assisted feedback, and client-side interaction into a unified mock interview platform."
    ],
}
for heading, paras in chapter5.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

add_table(
    doc,
    "Table 5.1 API Endpoints of the System",
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/health", "GET", "Verify backend availability"],
        ["/api/interview/resume", "POST", "Upload and parse resume"],
        ["/api/interview/questions", "POST", "Generate interview questions"],
        ["/api/interview/evaluate", "POST", "Evaluate interview responses"],
        ["/api/interview/ats-score", "POST", "Generate ATS analysis"],
    ],
)

add_table(
    doc,
    "Table 5.2 Evaluation Parameters",
    ["Parameter", "Description"],
    [
        ["Content Score", "Measures relevance and completeness of answer"],
        ["Speaking Score", "Approximates clarity and delivery quality"],
        ["Overall Score", "Combined answer performance indicator"],
        ["Strengths", "Positive aspects of answer quality"],
        ["Weaknesses", "Areas needing improvement"],
        ["Suggestions", "Recommended next-step improvements"],
    ],
)

add_table(
    doc,
    "Table 5.3 ATS Analysis Parameters",
    ["Parameter", "Description"],
    [
        ["Formatting", "Checks readability and parser-friendliness"],
        ["Section Completeness", "Checks presence of important resume sections"],
        ["Keyword Match", "Measures skill/job-term alignment"],
        ["Experience Quality", "Assesses action and impact representation"],
        ["Achievements Impact", "Checks quantified accomplishments"],
        ["Grammar Clarity", "Measures clarity and resume readability"],
    ],
)

for fig in [
    "[Insert Figure 5.1: Login Page]",
    "[Insert Figure 5.2: Resume Upload Interface]",
    "[Insert Figure 5.3: Interview Session Screen]",
    "[Insert Figure 5.4: ATS Analysis Dashboard]",
    "[Insert Figure 5.5: Result Summary Screen]",
]:
    add_paragraph(doc, fig, first_line_indent=False)

doc.add_page_break()

# Chapter 6
add_heading(doc, "CHAPTER 6", 1, center=True)
add_heading(doc, "CONCLUSION AND FUTURE SCOPE", 1, center=True)
chapter6 = {
    "6.1 Conclusion": [
        "The AI Interview System was developed to address an important and practical problem faced by students and job seekers: the lack of personalized, structured, and measurable interview preparation support. The project successfully combines multiple capabilities that are usually scattered across different tools, including resume parsing, ATS checking, interview simulation, answer evaluation, camera-based attentiveness feedback, and progress tracking.",
        "The implemented system demonstrates that a web application can meaningfully improve interview preparation by connecting the candidate's resume with relevant question generation and actionable answer feedback. Instead of merely presenting a static list of questions, the system provides a guided practice experience with interpretable results. In academic terms, the project validates the feasibility of integrating frontend engineering, backend services, AI interaction, dataset-backed logic, and basic computer vision into one coherent platform.",
        "The major objectives of the project were achieved. The application accepts resumes in multiple formats, extracts profile information, generates interview questions, evaluates responses, and computes ATS-oriented analysis. The use of gamification and local progress storage further strengthens the application's value as a continuous preparation tool rather than a one-time demonstration system."
    ],
    "6.2 Practical Usefulness": [
        "The practical value of the project lies in its relevance to real placement preparation. Students often need to prepare quickly, repeatedly, and independently. This system supports all three needs by allowing self-guided usage at any time. It can be especially useful before campus placements, internships, and early-career interviews where candidates need to practice resume explanation, technical confidence, and communication quality.",
        "By combining ATS analysis and interview practice, the system offers help at two important points in the hiring journey. This makes it more useful than tools that evaluate only answers or only resumes. Candidates can first improve their profile, then practice presenting that profile verbally."
    ],
    "6.3 Current Limitations": [
        "The current version has several limitations consistent with its academic scope. Authentication and persistence rely on local storage rather than secure multi-user server-side storage. Voice recognition quality depends on browser support and environmental noise. Camera monitoring provides attentiveness indicators rather than deeper emotion or behavior understanding. External AI dependency can also affect result depth when the service is unavailable or not configured.",
        "Another limitation is language scope. The system currently assumes English resume content and English interview flow. Evaluation is also optimized for general mock interview settings rather than industry-specific deep technical interviews. These limitations, however, create clear directions for future enhancement."
    ],
    "6.4 Future Scope": [
        "The future scope of the project is broad and promising. One direction is multilingual support so that users can practice interviews in languages beyond English. Another is real-time speech scoring using audio features such as pace, pauses, and confidence markers. Emotion analysis, facial expression understanding, and stronger anti-cheating signals could also make the camera mode more insightful.",
        "On the product side, cloud deployment, database-backed multi-user storage, recruiter dashboards, and institution-level analytics could transform the project from an academic prototype into a scalable service. Integration with placement cells, job portals, or skill recommendation systems could further increase practical impact. In summary, the current project establishes a strong foundation for future development in intelligent employability support systems."
    ],
}
for heading, paras in chapter6.items():
    add_heading(doc, heading, 2)
    for para in paras:
        add_paragraph(doc, para)

doc.add_page_break()
add_heading(doc, "REFERENCES", 1, center=True)
refs = [
    "[1] React Documentation, Meta Platforms, Inc. Available: https://react.dev/",
    "[2] Node.js Documentation. Available: https://nodejs.org/",
    "[3] Express.js Documentation. Available: https://expressjs.com/",
    "[4] MediaPipe Tasks Vision Documentation. Available: https://ai.google.dev/edge/mediapipe/",
    "[5] Mammoth Documentation for DOCX text extraction. Available: https://github.com/mwilliamson/mammoth.js",
    "[6] Applicant Tracking System concepts and resume screening studies, various online technical and academic sources.",
    "[7] Research articles on AI-assisted interview preparation and automated response evaluation.",
    "[8] OpenRouter API documentation. Available: https://openrouter.ai/docs",
    "[9] Documentation and references for browser-based speech recognition and web media device APIs.",
    "[10] Standard academic references related to human-computer interaction, automated assessment, and employability systems.",
]
for ref in refs:
    add_paragraph(doc, ref, first_line_indent=False, space_after=0)

doc.add_page_break()
add_heading(doc, "APPENDIX A", 1, center=True)
add_heading(doc, "SOURCE CODE OVERVIEW", 2)
appendix_a = [
    "The frontend codebase is organized around a React application structure with top-level routing and feature-oriented components. Authentication interfaces, dashboard features, shared UI helpers, and interview practice flow are defined in separate components to improve maintainability. Utility modules handle storage access, API requests, gamification behavior, and optional OpenRouter communication support.",
    "The backend codebase is centered on an Express server with endpoints for health checks, resume upload, question generation, answer evaluation, and ATS scoring. Business logic is implemented in a dedicated interview service module responsible for parsing documents, normalizing question banks, scoring ATS quality, evaluating answers, and shaping the final JSON payloads consumed by the frontend.",
    "For final submission, important code excerpts such as API route definitions, resume parsing logic, question selection logic, evaluation functions, ATS scoring logic, and camera monitoring integration can be attached as part of the appendix if required by the department."
]
for para in appendix_a:
    add_paragraph(doc, para)

doc.add_page_break()
add_heading(doc, "APPENDIX B", 1, center=True)
add_heading(doc, "PROJECT SCREENSHOTS", 2)
for shot in [
    "[Insert Screenshot: Login Page]",
    "[Insert Screenshot: Signup Page]",
    "[Insert Screenshot: Resume Upload Page]",
    "[Insert Screenshot: Interview Type Selection Page]",
    "[Insert Screenshot: Camera Interview Screen]",
    "[Insert Screenshot: Voice Interview Screen]",
    "[Insert Screenshot: ATS Analysis Page]",
    "[Insert Screenshot: Result Dashboard]",
    "[Insert Screenshot: Score Summary Page]",
]:
    add_paragraph(doc, shot, first_line_indent=False)

# footer page numbers
for sec in doc.sections:
    footer = sec.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(para)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
